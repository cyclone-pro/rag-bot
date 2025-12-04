"""Parse PDF, DOCX, TXT, and ZIP files."""
import importlib
import io
import logging
import zipfile
from typing import Tuple, List

logger = logging.getLogger(__name__)


def parse_file(filename: str, file_bytes: bytes) -> str:
    """
    Parse a resume file and extract text.
    
    Supports: PDF, DOCX, TXT, ZIP (containing supported files)
    
    Returns:
        Extracted text content
    """
    filename_lower = filename.lower()
    
    logger.info(f"   Parsing file type: {filename}")
    
    if filename_lower.endswith('.pdf'):
        logger.info("   → Detected: PDF")
        return _parse_pdf(file_bytes)
    elif filename_lower.endswith('.docx'):
        logger.info("   → Detected: DOCX")
        return _parse_docx(file_bytes)
    elif filename_lower.endswith('.doc'):
        logger.info("   → Detected: DOC (legacy)")
        return _parse_doc(file_bytes)
    elif filename_lower.endswith('.txt'):
        logger.info("   → Detected: TXT")
        text = file_bytes.decode('utf-8', errors='ignore')
        logger.info(f"   → Extracted {len(text)} chars from TXT")
        return text
    elif filename_lower.endswith('.zip'):
        logger.info("   → Detected: ZIP")
        return _parse_zip(file_bytes)
    else:
        logger.error(f"   ❌ Unsupported file format: {filename}")
        raise ValueError(f"Unsupported file format: {filename}")


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pypdf."""
    try:
        logger.info("   → Using pypdf library...")
        from pypdf import PdfReader
        
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        logger.info(f"   → PDF has {len(reader.pages)} pages")
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)
                logger.debug(f"   → Page {i+1}: {len(text)} chars")
        
        full_text = "\n".join(text_parts)
        logger.info(f"   → Extracted {len(full_text)} chars from PDF")
        return full_text
        
    except ImportError:
        logger.error("   ❌ pypdf not installed")
        raise ValueError("PDF parsing requires 'pypdf' library. Install: pip install pypdf")
    except Exception as e:
        logger.exception("   ❌ PDF parsing failed")
        raise ValueError(f"Failed to parse PDF: {e}")


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        logger.info("   → Using python-docx library...")
        from docx import Document
        
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        logger.info(f"   → DOCX has {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        logger.info(f"   → Extracted {len(full_text)} chars from DOCX")
        return full_text
        
    except ImportError:
        logger.error("   ❌ python-docx not installed")
        raise ValueError("DOCX parsing requires 'python-docx' library. Install: pip install python-docx")
    except Exception as e:
        logger.exception("   ❌ DOCX parsing failed")
        raise ValueError(f"Failed to parse DOCX: {e}")


def _parse_doc(file_bytes: bytes) -> str:
    """Extract text from legacy DOC using optional textract."""
    # textract 1.6.x currently requires pip<24.1 due to metadata issues; keep it optional.
    install_hint = (
        "DOC parsing requires optional 'textract'. "
        "Install with: pip install 'pip<24.1' && pip install textract==1.6.3, "
        "or convert the file to DOCX/PDF and retry."
    )
    try:
        logger.info("   → Using textract library...")
        textract = importlib.import_module("textract")
        text = textract.process(io.BytesIO(file_bytes)).decode('utf-8', errors='ignore')
        logger.info(f"   → Extracted {len(text)} chars from DOC")
        return text
    except ModuleNotFoundError:
        logger.error("   ❌ textract not installed")
        raise ValueError(install_hint)
    except Exception as e:
        logger.exception("   ❌ DOC parsing failed")
        raise ValueError(f"Failed to parse DOC: {e}. {install_hint}")


def _parse_zip(file_bytes: bytes) -> str:
    """Extract text from first supported file in ZIP."""
    try:
        logger.info("   → Opening ZIP archive...")
        zip_file = io.BytesIO(file_bytes)
        
        with zipfile.ZipFile(zip_file, 'r') as zf:
            members = zf.namelist()
            logger.info(f"   → ZIP contains {len(members)} files")
            
            # Find first supported file
            for member in members:
                if member.lower().endswith(('.pdf', '.docx', '.doc', '.txt')):
                    logger.info(f"   → Found resume in ZIP: {member}")
                    
                    member_bytes = zf.read(member)
                    return parse_file(member, member_bytes)
        
        logger.error("   ❌ No supported resume file found in ZIP")
        raise ValueError("No supported resume file (.pdf, .docx, .doc, .txt) found in ZIP")
        
    except zipfile.BadZipFile:
        logger.error("   ❌ Invalid ZIP file")
        raise ValueError("Invalid ZIP file")
    except Exception as e:
        logger.exception("   ❌ ZIP parsing failed")
        raise ValueError(f"Failed to parse ZIP: {e}")
